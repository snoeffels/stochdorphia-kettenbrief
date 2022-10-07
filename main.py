import os
import re
from pathlib import Path

import openpyxl
from docxcompose.composer import Composer
from docx import Document as Document_compose, Document

# Dateiname der Datenbank
excel = "example.xlsx"
excel_first_row_to_use = 2

# Dateiname der Vorlage
template = "template.docx"

# Dateiname des Resultats
output = "serienbrief.docx"

# Tempoärer Dateineime (ist egal aber sollte einzigartig sein)
temporary = "temporary.docx"

# Anreden mänlich weiblich
anrede_regex = r"%anrede%"
anrede_m = "Sehr geehrter Bundesbruder"
anrede_w = "Sehr geehrte Bundesschwester"


def merge(doc1, doc2, prepend_page_break):
    doc = Document_compose(doc1)

    if prepend_page_break:
        doc.add_page_break()

    composer = Composer(doc)
    composer.append(Document_compose(doc2))
    composer.save(doc1)


def replace_in_doc(doc, reg, rep):
    document = Document(doc)

    for p in document.paragraphs:
        if reg.search(p.text):
            inline = p.runs
            for x in range(len(inline)):
                if reg.search(inline[x].text):
                    text = reg.sub(rep, inline[x].text)
                    inline[x].text = text

    document.save(doc)


try:
    os.remove(output)
except:
    print("There is no " + output + " to delete")

Path(output).touch(exist_ok=True)

wb = openpyxl.load_workbook(excel)
sh = wb.active

for i in range(excel_first_row_to_use, sh.max_row + 1):
    print("row: " + str(i))
    Document(template).save(temporary)

    empty = False
    maskulin = False

    for j in range(1, 8):

        replace = r""
        if sh.cell(row=i, column=j).value is not None:

            value_string = str(sh.cell(row=i, column=j).value)

            if j == 1:
                maskulin = "Herr" in value_string
                replace = r"" + value_string
            elif j == 2:
                replace = r"" + value_string + " "
            elif j == 6:
                replace = r"" + str(int(sh.cell(row=i, column=j).value))
            else:
                replace = r"" + value_string

        else:
            if j == 1:
                try:
                    os.remove(temporary)
                except:
                    print("Could not delete " + temporary)
                print("Empty first column occurred")
                exit(0)
            else:
                print("None in row: " + str(i) + " col: " + str(j))

        replace_in_doc(temporary, re.compile(r"%" + str(j) + "%"), replace)

    if maskulin:
        replace_in_doc(temporary, re.compile(anrede_regex), anrede_m)
    else:
        replace_in_doc(temporary, re.compile(anrede_regex), anrede_w)

    if i == excel_first_row_to_use:
        Document(temporary).save(output)
    else:
        merge(output, temporary, True)
